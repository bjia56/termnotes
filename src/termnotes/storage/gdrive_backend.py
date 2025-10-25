"""
Google Drive-based note storage backend
"""

import os
import uuid
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Dict

from docx import Document

from ..utils import utc_now, normalize_to_utc
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaInMemoryUpload
from googleapiclient.errors import HttpError

from .base import StorageBackend
from ..note import Note


# Google Drive API scopes
SCOPES = ['https://www.googleapis.com/auth/drive.file']


class GoogleDriveBackend(StorageBackend):
    """Google Drive implementation of storage backend"""

    def __init__(
        self,
        credentials_path: Optional[str] = None,
        token_path: Optional[str] = None,
        app_folder: str = "termnotes"
    ):
        """
        Initialize Google Drive storage backend

        Args:
            credentials_path: Path to OAuth2 credentials JSON file
                            Default: ~/.termnotes/google_credentials.json
            token_path: Path to store user access token
                       Default: ~/.termnotes/google_token.json
            app_folder: Name of folder in Drive root to store notes
        """
        # Set up paths
        config_dir = Path.home() / ".termnotes"
        config_dir.mkdir(parents=True, exist_ok=True)

        self.credentials_path = credentials_path or str(config_dir / "google_credentials.json")
        self.token_path = token_path or str(config_dir / "google_token.json")
        self.app_folder_name = app_folder

        # Internal state
        self._service = None
        self._folder_id: Optional[str] = None
        self._file_id_map: Dict[str, str] = {}

        # Initialize
        self._authenticate()
        self._ensure_app_folder()

    def _authenticate(self):
        """Authenticate with Google Drive using OAuth2"""
        creds = None

        # Load existing token if available
        if os.path.exists(self.token_path):
            creds = Credentials.from_authorized_user_file(self.token_path, SCOPES)

        # If no valid credentials, need to authenticate
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                # Token expired - refresh it
                creds.refresh(Request())
            else:
                # No valid token - run OAuth flow
                if not os.path.exists(self.credentials_path):
                    raise FileNotFoundError(
                        f"OAuth credentials not found at {self.credentials_path}\n"
                        f"Please create OAuth2 credentials in Google Cloud Console:\n"
                        f"1. Go to https://console.cloud.google.com\n"
                        f"2. Create a project and enable Drive API\n"
                        f"3. Create OAuth2 credentials (Desktop app)\n"
                        f"4. Download and save to {self.credentials_path}"
                    )

                flow = InstalledAppFlow.from_client_secrets_file(
                    self.credentials_path,
                    SCOPES
                )
                creds = flow.run_local_server(port=0)

            # Save token for future runs
            with open(self.token_path, 'w') as token:
                token.write(creds.to_json())

            # Set restrictive permissions
            os.chmod(self.token_path, 0o600)

        # Build Drive API service
        self._service = build('drive', 'v3', credentials=creds)

    def _ensure_app_folder(self):
        """Ensure the termnotes folder exists in Drive root"""
        # Search for existing folder
        query = (
            f"name='{self.app_folder_name}' and "
            f"mimeType='application/vnd.google-apps.folder' and "
            f"trashed=false"
        )

        try:
            results = self._service.files().list(
                q=query,
                spaces='drive',
                fields='files(id, name)'
            ).execute()

            files = results.get('files', [])

            if files:
                # Folder exists - use first match
                self._folder_id = files[0]['id']
            else:
                # Create folder
                file_metadata = {
                    'name': self.app_folder_name,
                    'mimeType': 'application/vnd.google-apps.folder'
                }

                folder = self._service.files().create(
                    body=file_metadata,
                    fields='id'
                ).execute()

                self._folder_id = folder['id']

        except HttpError as e:
            raise Exception(f"Failed to access Google Drive: {e}")

    def _sync_file_id_map(self):
        """Build map of note_id -> drive_file_id from Drive"""
        query = (
            f"'{self._folder_id}' in parents and "
            f"name contains '.docx' and "
            f"trashed=false"
        )

        all_files = []
        page_token = None

        try:
            # Paginate through all results
            while True:
                results = self._service.files().list(
                    q=query,
                    spaces='drive',
                    fields='nextPageToken, files(id, name)',
                    pageSize=100,
                    pageToken=page_token
                ).execute()

                files = results.get('files', [])
                all_files.extend(files)

                page_token = results.get('nextPageToken')
                if not page_token:
                    break

            # Build map
            self._file_id_map.clear()
            for file in all_files:
                # Extract note_id from filename (remove .docx extension)
                filename = file['name']
                if filename.endswith('.docx'):
                    note_id = filename[:-5]
                    self._file_id_map[note_id] = file['id']

        except HttpError as e:
            raise Exception(f"Failed to list files from Drive: {e}")

    def get_all_notes(self) -> List[Note]:
        """Get all notes from Google Drive"""
        # Sync file list
        self._sync_file_id_map()

        notes = []

        for note_id in self._file_id_map.keys():
            note = self.get_note(note_id)
            if note:
                notes.append(note)

        # Sort by updated_at descending
        notes.sort(key=lambda n: n.updated_at, reverse=True)

        return notes

    def get_note(self, note_id: str) -> Optional[Note]:
        """Get a specific note by ID"""
        # Get Drive file ID
        file_id = self._file_id_map.get(note_id)

        if not file_id:
            # Not in cache - try to find it
            self._sync_file_id_map()
            file_id = self._file_id_map.get(note_id)

            if not file_id:
                return None

        # Download file content
        try:
            request = self._service.files().get_media(fileId=file_id)
            docx_bytes = request.execute()

            # Parse DOCX
            note = self._note_from_docx(note_id, docx_bytes)

            return note

        except HttpError as e:
            if e.resp.status == 404:
                # File deleted remotely - remove from cache
                self._file_id_map.pop(note_id, None)
                return None
            raise Exception(f"Failed to download note {note_id}: {e}")
        except Exception as e:
            raise Exception(f"Failed to parse note {note_id}: {e}")

    def save_note(self, note: Note):
        """Save or update a note in Google Drive"""
        # Update timestamp
        note.updated_at = utc_now()

        # Convert to DOCX
        docx_bytes = self._note_to_docx(note)

        # Check if file exists
        file_id = self._file_id_map.get(note.id)

        try:
            if file_id:
                # Update existing file
                media = MediaInMemoryUpload(
                    docx_bytes,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    resumable=True
                )

                self._service.files().update(
                    fileId=file_id,
                    media_body=media,
                    fields='id'
                ).execute()

            else:
                # Create new file
                file_metadata = {
                    'name': f'{note.id}.docx',
                    'parents': [self._folder_id],
                    'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                }

                media = MediaInMemoryUpload(
                    docx_bytes,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    resumable=True
                )

                new_file = self._service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id'
                ).execute()

                # Update cache
                self._file_id_map[note.id] = new_file['id']

        except HttpError as e:
            raise Exception(f"Failed to save note {note.id}: {e}")

    def delete_note(self, note_id: str):
        """Delete a note from Google Drive"""
        # Get file ID
        file_id = self._file_id_map.get(note_id)

        if not file_id:
            # Not found - might have been deleted already
            return

        try:
            # Delete from Drive
            self._service.files().delete(fileId=file_id).execute()

            # Remove from cache
            self._file_id_map.pop(note_id, None)

        except HttpError as e:
            if e.resp.status == 404:
                # Already deleted - just clean up cache
                self._file_id_map.pop(note_id, None)
            else:
                raise Exception(f"Failed to delete note {note_id}: {e}")

    def close(self):
        """Clean up resources"""
        self._file_id_map.clear()
        self._service = None

    def _note_to_docx(self, note: Note) -> bytes:
        """Convert Note object to DOCX format"""
        # Create new document
        doc = Document()

        # Add content as paragraphs
        if note.content:
            for line in note.content.split('\n'):
                doc.add_paragraph(line)

        # Set core properties for metadata
        doc.core_properties.title = note.id
        doc.core_properties.created = note.created_at
        doc.core_properties.modified = note.updated_at

        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read()

    def _note_from_docx(self, note_id: str, docx_bytes: bytes) -> Note:
        """Create Note object from DOCX bytes, preserving markdown format"""
        # Load document from bytes
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)

        # Extract metadata from core properties
        created_at = doc.core_properties.created or utc_now()
        updated_at = doc.core_properties.modified or utc_now()

        # Normalize datetimes to UTC (DOCX files may have timezone-aware timestamps)
        created_at = normalize_to_utc(created_at)
        updated_at = normalize_to_utc(updated_at)

        # Extract text directly from paragraphs to preserve markdown
        # Each paragraph in DOCX corresponds to a line in the original markdown
        paragraphs = [p.text for p in doc.paragraphs]
        content = '\n'.join(paragraphs)

        return Note(
            note_id=note_id,
            content=content,
            created_at=created_at,
            updated_at=updated_at
        )
